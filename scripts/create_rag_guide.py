from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor


OUT_DIR = Path("doc_output")
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "RAG_Implementation_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for margin in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin):
    pass
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for name, size, color in [("Title", 24, "000000"), ("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RAG Implementation Guide")
run.bold = True
run.font.size = Pt(24)
run.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("How the React + Express PDF chat app works end to end")
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor.from_string("555555")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Project: RAGReactApp | Scope: upload, parse, chunk, embed, retrieve, prompt, answer")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor.from_string("666666")

def add_heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    p.add_run(text)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)
    return p

def add_number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)
    return p

add_heading("1. What RAG Means")
add_body("RAG stands for Retrieval-Augmented Generation. The idea is simple: instead of asking the LLM to answer only from its built-in knowledge, we first retrieve relevant text from the user’s uploaded document, then pass that text into the model as context.")
add_body("In this project, the uploaded PDF is parsed into text, split into chunks, embedded into vectors, stored in a vector store, retrieved again when the user asks a question, and finally combined into a prompt that Gemini uses to generate the answer.")

add_heading("2. What This App Does")
add_bullet("Frontend: React + Vite chat interface with PDF upload, chat history, markdown rendering, loading state, and error handling.")
add_bullet("Backend: Express API for upload and chat, plus services for parsing, chunking, embeddings, vector storage, and RAG orchestration.")
add_bullet("AI: Gemini is used twice, once for embeddings and once for final answer generation.")
add_bullet("Storage: Pinecone is supported when configured, and a local in-memory fallback is used when Pinecone credentials are absent.")

add_heading("3. End-to-End Flow")
for step in [
    "User uploads a PDF from the frontend.",
    "The backend saves the file and parses it into text.",
    "The extracted text is chunked into smaller pieces with overlap.",
    "Each chunk is converted into a Gemini embedding.",
    "The embeddings are stored in Pinecone or local memory.",
    "When the user asks a question, the question is also embedded.",
    "The backend retrieves the most relevant chunks from vector storage.",
    "Those chunks are assembled into a context block.",
    "The context plus the question is sent to Gemini 2.5 Flash.",
    "Gemini produces the final answer, which is returned to the UI.",
]:
    add_number(step)

add_heading("4. File-by-File Implementation")
add_heading("Backend entry points", 2)
add_bullet("Server/src/server.js starts the Express app and listens on the configured port.")
add_bullet("Server/src/app.js wires middleware, CORS, JSON parsing, and the /upload and /chat routes.")
add_bullet("Server/src/routes/uploadRoutes.js handles PDF uploads with multer.")
add_bullet("Server/src/routes/chatRoutes.js handles user questions.")

add_heading("Parsing and upload", 2)
add_bullet("Server/src/middleware/uploadMiddleware.js configures multer disk storage and PDF-only validation.")
add_bullet("Server/src/controllers/uploadController.js receives the file, parses it, stores the latest document in memory, chunks it, and indexes it.")
add_bullet("Server/src/services/pdfService.js reads the PDF and extracts text using pdf-parse.")

add_heading("Chunking", 2)
add_bullet("Server/src/utils/chunkText.js splits document text into section-aware chunks.")
add_bullet("It detects common headings like Experience, Projects, Education, Skills, Summary, and keeps those sections grouped.")
add_bullet("Long sections are split again with overlap so context is preserved across chunk boundaries.")

add_heading("Embeddings and vectors", 2)
add_bullet("Server/src/services/embeddingService.js calls Gemini embedding model gemini-embedding-001.")
add_bullet("Server/src/services/vectorService.js stores vectors in Pinecone when configured.")
add_bullet("If Pinecone is not configured, it falls back to an in-memory vector store for local testing.")

add_heading("RAG orchestration", 2)
add_bullet("Server/src/services/ragService.js creates embeddings for each chunk during upload.")
add_bullet("It retrieves top matches for a user question and sorts them into a readable order.")
add_bullet("For broad document questions, it can use the full uploaded document context instead of only top-k chunks.")
add_bullet("It then prompts Gemini 2.5 Flash with the context and the question.")

add_heading("Frontend", 2)
add_bullet("client/src/pages/Home.jsx shows the document upload button and chat window.")
add_bullet("client/src/hooks/useChat.js manages chat state, upload state, errors, and document info.")
add_bullet("client/src/api/chatApi.js sends multipart upload requests and chat requests to the backend.")
add_bullet("client/src/components/ChatMessage.jsx renders markdown responses.")
add_bullet("client/src/components/ChatWindow.jsx auto-scrolls to the latest message.")

add_heading("5. How Chunking Works")
add_body("Chunking is important because long PDFs are too large to send directly to the embedding model or the chat model as one block. The chunker first normalizes the text, then looks for likely section headings. If it sees a heading, it starts a new section. That keeps important structure together, especially in resumes, reports, manuals, and letters.")
add_body("If a section is still too large, it is split into overlapping chunks. Overlap matters because key facts often sit near a boundary. Without overlap, one company name or one sentence can fall between chunks and be harder to retrieve later.")

add_heading("6. How Embeddings Work")
add_body("For each chunk, the backend sends the text to Gemini embedding model gemini-embedding-001. The model converts text into a numeric vector that represents semantic meaning. Similar ideas produce nearby vectors, which is what makes similarity search possible.")
add_body("The code for this is in Server/src/services/embeddingService.js. It calls ai.models.embedContent with contents: [text]. The response returns embedding values, which are stored along with the chunk metadata.")

add_heading("7. How Retrieval Works")
add_body("When the user asks a question, the backend embeds the question too. That question vector is compared against the stored document vectors. The vector store returns the top matches by semantic similarity.")
add_body("In local mode, the app computes cosine similarity in memory. In Pinecone mode, Pinecone handles the similarity search. In both cases, the goal is the same: find the document chunks that best match the question.")

add_heading("8. How Context Is Built")
add_body("The best matching chunks are not used individually. The app orders them by page and chunk index so the context reads naturally. It can also include neighboring chunks around the match to preserve surrounding meaning.")
add_body("For broad questions, the app may use the full uploaded text or the combined chunk set. This reduces omissions for questions like 'summarize this document' or 'list all key details.'")

add_heading("9. How the LLM Gets the Answer")
add_body("The final prompt includes instructions plus the retrieved context and the user question. Gemini 2.5 Flash then generates the response. The model is told not to invent details and to prefer direct extraction from the context.")
add_body("The answer is returned to the frontend and rendered in markdown. That is why formatted bullets, headings, and structure can appear nicely in the chat window.")

add_heading("10. Important Functions Used")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.autofit = False
table.columns[0].width = Inches(2.0)
table.columns[1].width = Inches(2.1)
table.columns[2].width = Inches(2.4)
hdr = table.rows[0].cells
hdr[0].text = "Function"
hdr[1].text = "File"
hdr[2].text = "Role"
for c in hdr:
    set_cell_shading(c, "E8EEF5")
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True

rows = [
    ("extractTextFromPdf", "pdfService.js", "Reads the PDF file and extracts plain text."),
    ("chunkText", "chunkText.js", "Splits PDF text into sections and overlapping chunks."),
    ("generateEmbedding", "embeddingService.js", "Converts text into vector embeddings using Gemini."),
    ("upsertVectors / queryVectors", "vectorService.js", "Stores and retrieves vectors from Pinecone or memory."),
    ("createEmbeddingsForDocument", "ragService.js", "Indexes every chunk from the uploaded PDF."),
    ("answerWithRag", "ragService.js", "Embeds the question, retrieves context, and calls the LLM."),
    ("uploadPDF", "uploadController.js", "Coordinates upload, parsing, storing, and indexing."),
]

for fn, file_name, role in rows:
    cells = table.add_row().cells
    cells[0].text = fn
    cells[1].text = file_name
    cells[2].text = role

add_heading("11. Why This Pipeline Is Reliable")
add_bullet("It preserves document structure better than raw chunk splitting.")
add_bullet("It stores both content and metadata, which improves retrieval order and traceability.")
add_bullet("It supports local mode without Pinecone, so development stays easy.")
add_bullet("It can switch to full-document context for broad questions, which improves completeness.")

add_heading("12. Current Limitations")
add_bullet("The local fallback is per server process and resets when the backend restarts.")
add_bullet("Very noisy OCR-style PDFs can still produce imperfect extracted text.")
add_bullet("If you want persistent production search, Pinecone should be configured.")

add_heading("13. Summary")
add_body("The app implements a practical RAG loop: upload, parse, chunk, embed, retrieve, assemble context, prompt Gemini, and return the response. The code is organized into small reusable services, so each part of the pipeline can be improved independently.")

doc.save(DOCX_PATH)
print(DOCX_PATH)
